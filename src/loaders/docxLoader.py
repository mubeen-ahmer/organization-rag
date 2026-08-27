import docx
from langchain_core.documents import Document

def loadDocx(path : str):
    doc = docx.Document(path)
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [Document(page_content=full_text, metadata={"source": path})]